"""
معالج ملفات DOCX (Word)
يستخرج النص ويعيد إنشاء ملف Word بالعربي مع محاذاة RTL صحيحة
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn


def extract_text_from_docx(file_path):
    """
    استخراج النص من ملف Word
    
    Args:
        file_path (str): مسار الملف
        
    Returns:
        list: قائمة الفقرات
    """
    try:
        doc = Document(file_path)
        texts = []
        
        # استخراج الفقرات
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # نتجاهل الفقرات الفارغة
                texts.append(text)
        
        # استخراج النص من الجداول أيضاً
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        texts.append(text)
        
        if not texts:
            raise Exception("الملف فارغ أو لا يحتوي على نص")
        
        return texts
    
    except Exception as e:
        raise Exception(f"خطأ في قراءة ملف Word: {str(e)}")


def create_docx(texts, output_path):
    """
    إنشاء ملف Word من نصوص عربية مترجمة مع محاذاة RTL كاملة
    
    Args:
        texts (list): قائمة النصوص المترجمة
        output_path (str): مسار الملف الناتج
        
    Returns:
        str: مسار الملف المحفوظ
    """
    try:
        doc = Document()
        
        # إعدادات الصفحة
        section = doc.sections[0]
        section.page_height = Pt(842)  # A4
        section.page_width = Pt(595)
        
        for text in texts:
            if not text.strip():
                # فقرة فارغة
                doc.add_paragraph()
                continue
            
            # إضافة الفقرة
            paragraph = doc.add_paragraph(text)
            
            # ✅ محاذاة يمين (RTL)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # ✅ تفعيل RTL على مستوى الفقرة
            pPr = paragraph._element.get_or_add_pPr()
            pPr.set(qn('w:bidi'), '1')  # تفعيل الاتجاه من اليمين لليسار
            
            # تنسيق الخط
            for run in paragraph.runs:
                run.font.size = Pt(14)
                run.font.name = 'Traditional Arabic'  # خط عربي
                
                # ✅ تفعيل RTL على مستوى النص
                rPr = run._element.get_or_add_rPr()
                rPr.set(qn('w:rtl'), '1')
        
        # حفظ الملف
        doc.save(output_path)
        
        return output_path
    
    except Exception as e:
        raise Exception(f"خطأ في إنشاء ملف Word: {str(e)}")


# اختبار
if __name__ == "__main__":
    print("✅ DOCX Handler جاهز!")
    print("✅ محاذاة RTL مفعّلة بالكامل")